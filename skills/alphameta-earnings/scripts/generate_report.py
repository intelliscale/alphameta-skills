"""
Earnings Update Report Generator
=================================
Single-file script for generating 8-12 page institutional-grade DOCX reports
with embedded matplotlib charts. Handles CJK/Latin bilingual fonts.

Usage (by LLM after collecting CLI data):
    from scripts.generate_report import DocxBuilder, ChartBuilder, analyst_pt_hbar
    cb = ChartBuilder()
    b = DocxBuilder(symbol="AAPL.US", company="Apple Inc.", ...)
    b.cover()
    b.section("1. Results Summary")
    b.table(headers, rows)
    b.image(cb.quarterly_bar(...))
    b.disclaimer()
    b.save()
"""

from __future__ import annotations

import io, os, tempfile
import numpy as np
import matplotlib
matplotlib.use("Agg")
from matplotlib import pyplot as plt, font_manager as fm
from matplotlib.figure import Figure
from matplotlib.axes import Axes
from matplotlib.lines import Line2D
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Constants ──
C_BLUE   = "#1A56DB"
C_ORANGE = "#F5A623"
C_GREEN  = "#27AE60"
C_RED    = "#E74C3C"
C_PURPLE = "#8E44AD"
C_GREY   = "#95A5A6"
_PALETTE = [C_BLUE, C_ORANGE, C_GREEN, C_RED, C_PURPLE, C_GREY]
DPI = 150

# ── CJK Font Setup ──
_CJK_FONTS = [
    'PingFang SC', 'Heiti SC', 'STHeiti', 'Songti SC',
    'SimHei', 'Microsoft YaHei',
    'Noto Sans CJK SC', 'WenQuanYi Micro Hei',
]
_available_fonts = {f.name for f in fm.fontManager.ttflist}
_CHART_FONT = next((f for f in _CJK_FONTS if f in _available_fonts), 'DejaVu Sans')
plt.rcParams['font.family'] = _CHART_FONT
plt.rcParams['axes.unicode_minus'] = False

# ── DOCX Constants ──
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x37, 0x40, 0x51)
GREY_RGB = RGBColor(0x6B, 0x72, 0x80)
HEAD_FILL = "DBE9FF"
DEFAULT_LATIN = "Calibri"
DEFAULT_CJK   = "Microsoft YaHei"


# ═══════════════════════════════════════════════════════════
#  CHART FUNCTIONS
# ═══════════════════════════════════════════════════════════

def _savefig(fig: Figure, prefix: str = "chart") -> str:
    fd, path = tempfile.mkstemp(prefix=f"{prefix}_", suffix=".png", dir="/tmp")
    os.close(fd)
    fig.savefig(path, dpi=DPI, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _style_ax(ax: Axes, ylabel: str = "", xlabel: str = "") -> None:
    ax.set_facecolor("white")
    ax.grid(axis="y", color="#DDDDDD", linewidth=0.8, zorder=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#CCCCCC")
    ax.spines["bottom"].set_color("#CCCCCC")
    if ylabel: ax.set_ylabel(ylabel, fontsize=10)
    if xlabel: ax.set_xlabel(xlabel, fontsize=10)


def quarterly_bar(title, quarters, values, estimate_idx=-1, ylabel="Revenue",
                  color=C_BLUE, est_color=C_ORANGE, figsize=(10, 5)):
    fig, ax = plt.subplots(figsize=figsize)
    colors = [est_color if i == (len(values) + estimate_idx if estimate_idx < 0 else estimate_idx)
              else color for i in range(len(values))]
    bars = ax.bar(quarters, values, color=colors, width=0.6, zorder=3)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(values)*0.01,
                f"{val:g}", ha="center", va="bottom", fontsize=9, fontweight="bold")
    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    ax.set_ylim(0, max(values)*1.15)
    ax.tick_params(axis="x", rotation=30, labelsize=9)
    _style_ax(ax, ylabel=ylabel)
    fig.tight_layout()
    return _savefig(fig, "quarterly_bar")


def growth_lines(title, quarters, series, ylabel="YoY Growth (%)", figsize=(10, 5)):
    fig, ax = plt.subplots(figsize=figsize)
    for s in series:
        ax.plot(quarters, s["values"], color=s.get("color", C_BLUE),
                marker=s.get("marker", "o"), linewidth=2, markersize=7, label=s["label"], zorder=3)
        for x, y in zip(quarters, s["values"]):
            ax.annotate(f"{y:g}%", (x, y), textcoords="offset points", xytext=(0, 8),
                        ha="center", fontsize=8, color=s.get("color", C_BLUE), fontweight="bold")
    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    ax.legend(fontsize=9, framealpha=0.9)
    ax.tick_params(axis="x", rotation=20, labelsize=9)
    ax.set_ylim(0, max(max(s["values"]) for s in series)*1.25)
    _style_ax(ax, ylabel=ylabel)
    fig.tight_layout()
    return _savefig(fig, "growth_lines")


def grouped_bar(title, categories, series, ylabel="", figsize=(11, 5)):
    n_cats, n_series = len(categories), len(series)
    x = np.arange(n_cats)
    width = 0.35 if n_series == 2 else 0.8/n_series
    fig, ax = plt.subplots(figsize=figsize)
    offsets = np.linspace(-(n_series-1)/2*width, (n_series-1)/2*width, n_series)
    for s, offset in zip(series, offsets):
        bars = ax.bar(x+offset, s["values"], width*0.95, label=s["label"],
                      color=s.get("color", C_BLUE), zorder=3)
        for bar, val in zip(bars, s["values"]):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+max(s["values"])*0.01,
                    f"{val:g}", ha="center", va="bottom", fontsize=8)
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=9)
    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    ax.legend(fontsize=9)
    all_vals = [v for s in series for v in s["values"]]
    ax.set_ylim(0, max(all_vals)*1.18)
    _style_ax(ax, ylabel=ylabel)
    fig.tight_layout()
    return _savefig(fig, "grouped_bar")


def pie_pair(title1, title2, labels, values1, values2, figsize=(12, 5)):
    colors = _PALETTE[:len(labels)]
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=figsize)
    def _pie(ax, vals, title):
        wedges, texts, autotexts = ax.pie(vals, labels=None, autopct="%1.0f%%",
                colors=colors, startangle=90, pctdistance=0.75)
        for at in autotexts:
            at.set_fontsize(10); at.set_fontweight("bold")
        ax.set_title(title, fontsize=13, fontweight="bold", pad=8)
    _pie(ax1, values1, title1); _pie(ax2, values2, title2)
    patches = [mpatches.Patch(color=c, label=l) for c, l in zip(colors, labels)]
    fig.legend(handles=patches, loc="lower center", ncol=min(len(labels), 3),
               fontsize=9, bbox_to_anchor=(0.5, -0.05))
    fig.tight_layout()
    return _savefig(fig, "pie_pair")


def scenario_hbar(title, scenarios, values, current, current_label, colors=None,
                  xlabel="Implied Price", figsize=(10, 4)):
    if colors is None: colors = [C_GREEN, C_BLUE, C_RED]
    fig, ax = plt.subplots(figsize=figsize)
    y = np.arange(len(scenarios))
    bars = ax.barh(y, values, color=colors, height=0.5, zorder=3)
    for bar, val, scenario in zip(bars, values, scenarios):
        pct = (val/current-1)*100
        label = f"${val:,.0f} ({'+' if pct>=0 else ''}{pct:.0f}%)"
        ax.text(bar.get_width()+max(values)*0.01, bar.get_y()+bar.get_height()/2,
                label, va="center", fontsize=10, fontweight="bold")
    ax.axvline(current, color=C_RED, linestyle="--", linewidth=1.5, zorder=4)
    ax.set_yticks(y); ax.set_yticklabels(scenarios, fontsize=11)
    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    ax.set_xlim(0, max(values)*1.25)
    ax.grid(axis="x", color="#DDDDDD", linewidth=0.8, zorder=0)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.legend(handles=[Line2D([0],[0],color=C_RED,linestyle="--",label=f"Current: {current_label}")],
              fontsize=9, loc="lower right")
    _style_ax(ax, xlabel=xlabel)
    fig.tight_layout()
    return _savefig(fig, "scenario_hbar")


def peer_multiples(title_left, title_right, companies, vals_left, vals_right,
                   highlight, xlabel_left="P/E (TTM)", xlabel_right="P/B", figsize=(13, 5)):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=figsize)
    def _hbar(ax, vals, title, xlabel):
        clean = [(c, v) for c, v in zip(companies, vals) if v is not None]
        names = [x[0] for x in clean]; numbers = [x[1] for x in clean]
        colors = [C_RED if n == highlight else C_BLUE for n in names]
        y = np.arange(len(names))
        bars = ax.barh(y, numbers, color=colors, height=0.5, zorder=3)
        for bar, val in zip(bars, numbers):
            ax.text(bar.get_width()+max(numbers)*0.02, bar.get_y()+bar.get_height()/2,
                    f"{val:.1f}x", va="center", fontsize=9)
        ax.set_yticks(y); ax.set_yticklabels(names, fontsize=10)
        ax.set_title(title, fontsize=12, fontweight="bold")
        ax.set_xlim(0, max(numbers)*1.3)
        ax.grid(axis="x", color="#DDDDDD", linewidth=0.8, zorder=0)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        ax.set_xlabel(xlabel, fontsize=9)
    _hbar(ax1, vals_left, title_left, xlabel_left)
    _hbar(ax2, vals_right, title_right, xlabel_right)
    fig.tight_layout(pad=3)
    return _savefig(fig, "peer_multiples")


def price_action(title, dates, prices, current, current_label,
                 target=None, target_label="Consensus PT", ylabel="Price", figsize=(12, 4.5)):
    fig, ax = plt.subplots(figsize=figsize)
    x = np.arange(len(dates))
    ax.plot(x, prices, color=C_BLUE, linewidth=2, zorder=3)
    ax.fill_between(x, prices, alpha=0.15, color=C_BLUE)
    ax.axhline(current, color=C_RED, linestyle="--", linewidth=1.2, zorder=4)
    legend = [Line2D([0],[0],color=C_RED,linestyle="--",label=f"Current: {current_label}")]
    if target is not None:
        ax.axhline(target, color=C_GREEN, linestyle="--", linewidth=1.2, zorder=4)
        legend.append(Line2D([0],[0],color=C_GREEN,linestyle="--",label=target_label))
    step = max(1, len(dates)//10)
    ax.set_xticks(x[::step]); ax.set_xticklabels(dates[::step], rotation=30, fontsize=8)
    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    ax.legend(handles=legend, fontsize=9)
    _style_ax(ax, ylabel=ylabel)
    fig.tight_layout()
    return _savefig(fig, "price_action")


def analyst_pt_hbar(title, current, current_label="Current",
                    target_mean=None, target_high=None, target_low=None,
                    num_analysts=None, xlabel="Price", figsize=(10, 3.5)):
    if target_mean is None:
        fig, ax = plt.subplots(figsize=(8, 2))
        ax.text(0.5, 0.5, "No consensus target data", ha="center", va="center",
                fontsize=12, color=C_GREY, transform=ax.transAxes)
        ax.set_title(title, fontsize=13, fontweight="bold"); ax.axis("off")
        fig.tight_layout(); return _savefig(fig, "analyst_pt_hbar")
    lo = target_low if target_low is not None else target_mean*0.8
    hi = target_high if target_high is not None else target_mean*1.2
    margin = (hi-lo)*0.25
    fig, ax = plt.subplots(figsize=figsize)
    ax.axvspan(lo, hi, alpha=0.06, color=C_BLUE, zorder=1)
    ax.plot([lo, hi], [0, 0], color="#B0C4DE", linewidth=5, zorder=2, solid_capstyle="butt")
    ax.scatter([lo, target_mean, hi], [0, 0, 0], color=[C_ORANGE, C_BLUE, C_GREEN],
               s=[100, 220, 100], zorder=6, edgecolors="white", linewidth=2)
    ax.text(lo, -0.8, f"Low ${lo:,.0f}", ha="center", va="top", fontsize=9, color=C_ORANGE, fontweight="bold")
    ax.text(target_mean, 0.8, f"Target ${target_mean:,.0f}", ha="center", va="bottom",
            fontsize=12, color=C_BLUE, fontweight="bold")
    ax.text(hi, -0.8, f"High ${hi:,.0f}", ha="center", va="top", fontsize=9, color=C_GREEN, fontweight="bold")
    ax.axvline(current, color=C_RED, linestyle="--", linewidth=1.5, zorder=4)
    pct = (current/target_mean-1)*100
    ax.text(current, 1.8, f"{'▼' if pct<0 else '▲'} {current_label} ({pct:+.1f}%)",
            ha="center", va="bottom", fontsize=9, color=C_RED, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.25", facecolor="#FFF0F0", edgecolor=C_RED, alpha=0.85))
    subtitle = f"based on {num_analysts} analysts" if num_analysts else "consensus targets"
    ax.set_title(f"{title}\n({subtitle})", fontsize=13, fontweight="bold", pad=15)
    ax.set_yticks([]); ax.set_ylim(-2.5, 3.0)
    ax.set_xlim(min(lo, current)-margin, max(hi, current)+margin)
    ax.set_xlabel(xlabel, fontsize=9)
    ax.grid(axis="x", color="#DDDDDD", linewidth=0.8, zorder=0)
    for sp in ["top","right","left"]: ax.spines[sp].set_visible(False)
    ax.tick_params(left=False)
    fig.tight_layout(); return _savefig(fig, "analyst_pt_hbar")


class ChartBuilder:
    quarterly_bar = staticmethod(quarterly_bar)
    growth_lines = staticmethod(growth_lines)
    grouped_bar = staticmethod(grouped_bar)
    pie_pair = staticmethod(pie_pair)
    scenario_hbar = staticmethod(scenario_hbar)
    peer_multiples = staticmethod(peer_multiples)
    price_action = staticmethod(price_action)
    analyst_pt_hbar = staticmethod(analyst_pt_hbar)


# ═══════════════════════════════════════════════════════════
#  DOCX BUILDER
# ═══════════════════════════════════════════════════════════

def _set_run_fonts(run, latin: str, cjk: str) -> None:
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), cjk)


def _set_doc_default_cjk(doc: Document, cjk: str = DEFAULT_CJK) -> None:
    for name in ("Normal", "Default Paragraph Font"):
        try:
            rPr = doc.styles[name].element.get_or_add_rPr()
            rPr.get_or_add_rFonts().set(qn("w:eastAsia"), cjk)
        except Exception:
            pass


class DocxBuilder:
    def __init__(self, symbol: str, company: str, report_date: str,
                 analysis_date: str, price: str, market_cap: str,
                 valuation: str, rating: str, output_path: str,
                 lang: str = "en", latin: str = DEFAULT_LATIN, cjk: str = DEFAULT_CJK):
        self.symbol = symbol; self.company = company
        self.report_date = report_date; self.analysis_date = analysis_date
        self.price = price; self.market_cap = market_cap
        self.valuation = valuation; self.rating = rating
        self.output_path = output_path; self.lang = lang
        self.latin = latin; self.cjk = cjk
        self.doc = Document()
        sec = self.doc.sections[0]
        sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.2); sec.right_margin = Inches(1.2)
        _set_doc_default_cjk(self.doc, cjk)

    def _run(self, para, text, bold=False, size=10, color=None, latin=None, cjk=None):
        run = para.add_run(str(text))
        run.bold = bold; run.font.size = Pt(size)
        _set_run_fonts(run, latin or self.latin, cjk or self.cjk)
        if color: run.font.color.rgb = color
        return run

    def page_break(self): self.doc.add_page_break()
    def cover(self):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(48)
        self._run(p, f"{self.company}（{self.symbol}）", bold=True, size=22, color=BLUE)
        p2 = self.doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p2, "Earnings Update", bold=True, size=16)
        self.doc.add_paragraph()
        p3 = self.doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p3, f"Analysis: {self.analysis_date}  |  Report: {self.report_date}", size=11, color=GREY_RGB)
        self.doc.add_paragraph()
        p4 = self.doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p4, f"Price: {self.price}  |  Mkt Cap: {self.market_cap}  |  "
                      f"{self.valuation}  |  {self.rating}", size=10)
        self.doc.add_page_break()

    def toc(self, sections):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
        self._run(p, "Contents", bold=True, size=15, color=BLUE)
        for num, title in sections:
            row = self.doc.add_paragraph()
            row.paragraph_format.space_before = Pt(3); row.paragraph_format.space_after = Pt(3)
            self._run(row, num+"  ", bold=True, size=10, color=BLUE)
            self._run(row, title, size=10)
        self.doc.add_page_break()

    def section(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
        self._run(p, text, bold=True, size=15, color=BLUE)

    def subsection(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
        self._run(p, text, bold=True, size=12, color=DARK)

    def body(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
        self._run(p, text, size=10)

    def bullet(self, text):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
        self._run(p, text, size=10)

    def qa(self, question, answer, watch):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        self._run(p, question, bold=True, size=10)
        for label, val in [("  ↳ Response: ", answer), ("  ↳ Watch: ", watch)]:
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.3)
            p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(3)
            self._run(p2, label, bold=True, size=9, color=DARK)
            self._run(p2, val, size=9)

    def table(self, headers, rows, col_widths=None):
        t = self.doc.add_table(rows=1+len(rows), cols=len(headers))
        t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]; cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(h)); run.bold = True; run.font.size = Pt(9)
            _set_run_fonts(run, self.latin, self.cjk)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"), HEAD_FILL); cell._tc.get_or_add_tcPr().append(shd)
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = t.rows[r_i+1].cells[c_i]; cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(str(val)); run.font.size = Pt(9)
                _set_run_fonts(run, self.latin, self.cjk)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows: row.cells[i].width = Inches(w)
        self.doc.add_paragraph()

    def image(self, path, width=6.0, caption=""):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(); run.add_picture(path, width=Inches(width))
        if caption:
            cp = self.doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(0); cp.paragraph_format.space_after = Pt(6)
            self._run(cp, caption, size=8, color=GREY_RGB)
        else: self.doc.add_paragraph()

    def disclaimer(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"),"single")
        bottom.set(qn("w:sz"),"6"); bottom.set(qn("w:space"),"1")
        bottom.set(qn("w:color"),"C0C0C0")
        pBdr.append(bottom); pPr.append(pBdr)
        self.subsection("Disclaimer")
        self.body("This report is for informational and educational purposes only. "
                  "Data sources include AlphaMeta API, SEC filings, and public information. "
                  "All estimates and forecasts involve material uncertainty. "
                  "Past performance does not guarantee future results.")

    def save(self) -> str:
        self.doc.save(self.output_path)
        return self.output_path
